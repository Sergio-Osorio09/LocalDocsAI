"""Shared pytest fixtures that generate synthetic test documents on the fly."""

from __future__ import annotations

from pathlib import Path

import fitz
import pytest
from docx import Document as DocxDocument
from openpyxl import Workbook


@pytest.fixture()
def sample_pdf(tmp_path: Path) -> Path:
    """Minimal PDF with two pages and an embedded norm code."""
    pdf_path = tmp_path / "Osinergmin-029-2016-OS-CD.pdf"
    doc = fitz.open()

    page1 = doc.new_page()
    page1.insert_text((50, 50), "Encabezado repetido", fontsize=10)
    page1.insert_text(
        (50, 150),
        "RCD N° 029-2016-OS/CD — Reglamento de distribución de gas natural",
        fontsize=12,
    )
    page1.insert_text(
        (50, 200),
        "Artículo 1°. Objeto. La presente norma establece...",
        fontsize=11,
    )
    page1.insert_text((50, 750), "Página 1", fontsize=9)

    page2 = doc.new_page()
    page2.insert_text((50, 50), "Encabezado repetido", fontsize=10)
    page2.insert_text(
        (50, 150),
        "Artículo 2°. Definiciones. Para efectos de esta norma...",
        fontsize=11,
    )
    page2.insert_text((50, 750), "Página 2", fontsize=9)

    doc.save(str(pdf_path))
    doc.close()
    return pdf_path


@pytest.fixture()
def sample_docx(tmp_path: Path) -> Path:
    """Minimal DOCX with a heading, paragraphs, and a table."""
    docx_path = tmp_path / "informe_test.docx"
    document = DocxDocument()

    document.add_heading("Informe de prueba", level=1)
    document.add_paragraph(
        "D.S. N° 021-2021-MINEM-EM establece los requisitos técnicos para instalaciones."
    )
    document.add_paragraph("Este es un párrafo de contenido normal.")

    table = document.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "Norma"
    table.cell(0, 1).text = "Año"
    table.cell(0, 2).text = "Estado"
    table.cell(1, 0).text = "RCD N° 029-2016-OS/CD"
    table.cell(1, 1).text = "2016"
    table.cell(1, 2).text = "Vigente"

    document.save(str(docx_path))
    return docx_path


@pytest.fixture()
def sample_xlsx(tmp_path: Path) -> Path:
    """Minimal XLSX with two sheets."""
    xlsx_path = tmp_path / "indicadores_test.xlsx"
    wb = Workbook()

    ws1 = wb.active
    ws1.title = "Hoja1"
    ws1.append(["Norma", "Descripción", "Código"])
    ws1.append(["Osinergmin-029-2016-OS-CD", "Distribución GN", "DS-001"])

    ws2 = wb.create_sheet("Hoja2")
    ws2.append(["Indicador", "Valor"])
    ws2.append(["Cobertura", "85%"])

    wb.save(str(xlsx_path))
    return xlsx_path
