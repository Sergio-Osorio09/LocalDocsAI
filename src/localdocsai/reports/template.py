"""Template loading, placeholder detection, and filling for Word reports."""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document

# Matches any [text inside brackets] — greedy up to first ]
_BRACKET_RE = re.compile(r"\[([^\]]+)\]")

# Inline bold pattern *Label:* at paragraph start
_BOLD_LABEL_RE = re.compile(r"^\*([^*]+)\*\s*:?\s*")

# Section heading pattern ## N. Title or # Title
_HEADING_RE = re.compile(r"^#{1,3}\s*\d*\.?\s*(.+)")

# Semantic type keywords (lowercase)
_SEMANTIC_MAP = {
    "date": ("fecha", "date"),
    "title": ("proyecto", "referencia", "nombre del proyecto"),
    "topic": ("tema", "área", "area", "asunto"),
    "objective": ("objetivo", "propósito", "proposito", "consulta", "breve descripción"),
    "synthesis": (
        "síntesis", "sintesis", "respuesta", "hallazgo", "información", "resumen",
        "aquí pegas", "puntos clave",
    ),
}


@dataclass
class TemplateField:
    """One replaceable field detected in a Word template."""

    paragraph_idx: int
    pattern: str      # full bracketed text e.g. "[DD/MM/AAAA]"
    context: str      # paragraph text without the placeholder (for labeling)
    label: str        # human-readable field name shown in the editor
    semantic_type: str  # "date" | "title" | "topic" | "objective" | "synthesis" | "unknown"
    value: str = ""   # filled by generator or user


def detect_fields(doc_path: Path) -> list[TemplateField]:
    """Return one TemplateField per bracketed placeholder found in *doc_path*."""
    doc = Document(str(doc_path))
    para_texts = ["".join(r.text for r in p.runs) for p in doc.paragraphs]
    fields: list[TemplateField] = []

    for idx, full_text in enumerate(para_texts):
        m = _BRACKET_RE.search(full_text)
        if not m:
            continue
        pattern = m.group(0)
        context = full_text.replace(pattern, "").strip()

        # If the placeholder is the whole paragraph, look back for the nearest heading
        if not context:
            context = _find_preceding_heading(para_texts, idx)

        label = _extract_label(context, f"Campo {len(fields) + 1}")
        sem = _classify(context + " " + label)
        fields.append(
            TemplateField(
                paragraph_idx=idx,
                pattern=pattern,
                context=context,
                label=label,
                semantic_type=sem,
            )
        )

    return fields


def fill_and_save(doc_path: Path, fields: list[TemplateField], output_path: Path) -> None:
    """Replace each field's placeholder with its value and write to *output_path*."""
    doc = Document(str(doc_path))
    for field in fields:
        para = doc.paragraphs[field.paragraph_idx]
        _replace_in_paragraph(para, field.pattern, field.value)
    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------


def _find_preceding_heading(para_texts: list[str], current_idx: int) -> str:
    """Walk backwards from *current_idx* to find the nearest non-empty heading paragraph."""
    for i in range(current_idx - 1, -1, -1):
        text = para_texts[i].strip()
        if text and not _BRACKET_RE.search(text):
            return text
    return ""


def _extract_label(context: str, fallback: str) -> str:
    """Derive a short label from the paragraph context text."""
    m = _BOLD_LABEL_RE.match(context)
    if m:
        return m.group(1).rstrip(":").strip()
    m = _HEADING_RE.match(context)
    if m:
        return m.group(1).strip()
    # Strip markdown noise and truncate
    cleaned = re.sub(r"[*#\-_]", "", context).strip()
    return cleaned[:60] if cleaned else fallback[:60]


def _classify(text: str) -> str:
    lower = text.lower()
    for sem_type, keywords in _SEMANTIC_MAP.items():
        if any(kw in lower for kw in keywords):
            return sem_type
    return "unknown"


def _replace_in_paragraph(para: object, pattern: str, replacement: str) -> None:
    """Replace *pattern* with *replacement* in a paragraph, keeping the first run's style."""
    runs = para.runs  # type: ignore[attr-defined]
    full_text = "".join(r.text for r in runs)
    if pattern not in full_text:
        return
    new_text = full_text.replace(pattern, replacement)
    if runs:
        runs[0].text = new_text
        for run in runs[1:]:
            run.text = ""
