"""High-level report writer — generates .docx from a filled template or plain conversation."""

from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path
from typing import Any

from localdocsai.reports.template import TemplateField, fill_and_save

# Brand palette — kept in sync with src/localdocsai/ui/themes/dark.qss
_ACCENT_HEX = (0x4E, 0xC2, 0xE8)  # cyan-blue (AI accent)
_CITE_HEX = (0xB8, 0x86, 0x1B)  # darker amber for print contrast
_TEXT_HEX = (0x1F, 0x23, 0x29)  # near-black body text
_MUTED_HEX = (0x6E, 0x73, 0x7D)  # gray for meta
_PANEL_HEX = (0xF4, 0xF6, 0xFA)  # very light cool gray for boxes
_CITE_RE = re.compile(r"\[(\d+)\]")


def write_docx(
    template_path: Path,
    fields: list[TemplateField],
    output_path: Path,
) -> None:
    """Fill *template_path* with *fields* values and save to *output_path*."""
    fill_and_save(template_path, fields, output_path)


def write_conversation_pdf(
    messages: list[dict[str, Any]],
    output_path: Path,
    title: str = "Conversación",
    *,
    model_name: str = "",
    app_version: str = "v0.1 · OFFLINE",
) -> None:
    """Export *messages* as a PDF that mirrors the docx layout.

    Uses reportlab so no external converter (LibreOffice, Word COM) is
    required. The visual structure matches write_conversation_docx:
    title, metadata table, then for every Q/A pair a numbered heading,
    the question in a tinted box, the response with amber [N] chips,
    and a detailed source list.
    """
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        HRFlowable,
        PageBreak,  # noqa: F401  (kept for future use)
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    def hex_rgb(rgb: tuple[int, int, int]) -> colors.Color:
        return colors.Color(rgb[0] / 255, rgb[1] / 255, rgb[2] / 255)

    text_col = hex_rgb(_TEXT_HEX)
    muted_col = hex_rgb(_MUTED_HEX)
    cite_col = hex_rgb(_CITE_HEX)
    panel_col = hex_rgb(_PANEL_HEX)

    styles = getSampleStyleSheet()
    base_font = "Helvetica"

    title_style = ParagraphStyle(
        "TitleStyle", parent=styles["Title"], fontName=f"{base_font}-Bold",
        fontSize=20, leading=24, textColor=text_col, spaceAfter=4, alignment=0,
    )
    subtitle_style = ParagraphStyle(
        "Subtitle", parent=styles["Normal"], fontName=f"{base_font}-Oblique",
        fontSize=10, textColor=muted_col, spaceAfter=12,
    )
    h2_style = ParagraphStyle(
        "H2", parent=styles["Heading2"], fontName=f"{base_font}-Bold",
        fontSize=14, leading=18, textColor=text_col, spaceBefore=16, spaceAfter=4,
    )
    label_style = ParagraphStyle(
        "Label", parent=styles["Normal"], fontName=f"{base_font}-Bold",
        fontSize=8, textColor=muted_col, spaceBefore=4, spaceAfter=2,
    )
    body_style = ParagraphStyle(
        "Body", parent=styles["Normal"], fontName=base_font,
        fontSize=11, leading=15, textColor=text_col, spaceAfter=4,
    )
    boxed_body_style = ParagraphStyle(
        "BoxedBody", parent=body_style, backColor=panel_col,
        borderPadding=8, leftIndent=4, rightIndent=4, spaceAfter=10,
    )
    source_head_style = ParagraphStyle(
        "SrcHead", parent=body_style, leftIndent=0.4 * cm, spaceBefore=4, spaceAfter=0,
    )
    snippet_style = ParagraphStyle(
        "Snippet", parent=styles["Italic"], fontName=f"{base_font}-Oblique",
        fontSize=10, textColor=muted_col, leftIndent=1.0 * cm, rightIndent=0.4 * cm,
        leading=13, spaceAfter=0,
    )
    chunkid_style = ParagraphStyle(
        "ChunkId", parent=body_style, fontName="Courier", fontSize=8,
        textColor=muted_col, leftIndent=1.0 * cm, spaceBefore=0, spaceAfter=6,
    )
    meta_label_style = ParagraphStyle(
        "MetaLabel", parent=body_style, fontName=f"{base_font}-Bold",
        fontSize=8, textColor=muted_col,
    )
    meta_value_style = ParagraphStyle(
        "MetaValue", parent=body_style, fontSize=10, textColor=text_col,
    )

    def cite_color_hex() -> str:
        return f"#{_CITE_HEX[0]:02x}{_CITE_HEX[1]:02x}{_CITE_HEX[2]:02x}"

    def format_inline(text: str) -> str:
        """Convert plain text with [N] markers to reportlab markup.

        Citation tokens render as bold colored monospace; the rest is
        escaped so stray < > & in the answer do not break parsing.
        """
        cite_hex = cite_color_hex()
        parts: list[str] = []
        last = 0
        for m in _CITE_RE.finditer(text):
            if m.start() > last:
                parts.append(_xml_escape(text[last : m.start()]))
            parts.append(
                f'<font face="Courier-Bold" color="{cite_hex}" size="10">'
                f"{m.group(0)}</font>"
            )
            last = m.end()
        if last < len(text):
            parts.append(_xml_escape(text[last:]))
        return "".join(parts)

    story: list[Any] = []
    story.append(Paragraph(_xml_escape(title), title_style))
    story.append(Paragraph("Conversación exportada de LocalDocsAI", subtitle_style))

    # Metadata table
    now = datetime.now()
    total_sources = sum(
        len(m.get("sources", [])) for m in messages if m.get("role") == "assistant"
    )
    questions = sum(1 for m in messages if m.get("role") == "user")
    meta_rows = [
        ("FECHA", now.strftime("%d/%m/%Y · %H:%M")),
        ("APLICACIÓN", f"LocalDocsAI  {app_version}"),
        ("MODELO", model_name or "—"),
        ("PREGUNTAS", str(questions)),
        ("FUENTES CITADAS", str(total_sources)),
    ]
    table_data = [
        [Paragraph(k, meta_label_style), Paragraph(_xml_escape(v), meta_value_style)]
        for k, v in meta_rows
    ]
    meta_table = Table(table_data, colWidths=[4.5 * cm, 12.0 * cm])
    meta_table.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LINEBELOW", (0, 0), (-1, -1), 0.4, colors.Color(0.87, 0.89, 0.92)),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    story.append(meta_table)
    story.append(Spacer(1, 0.4 * cm))
    story.append(HRFlowable(width="100%", thickness=0.6, color=colors.Color(0.87, 0.89, 0.92)))
    story.append(Spacer(1, 0.2 * cm))

    # Conversation
    question_index = 0
    for msg in messages:
        role = msg.get("role", "")
        text = msg.get("text", "")
        if role == "user":
            question_index += 1
            short = text.strip().split("\n")[0][:80] or "Pregunta"
            story.append(Paragraph(f"{question_index}.  {_xml_escape(short)}", h2_style))
            story.append(Paragraph("PREGUNTA", label_style))
            story.append(Paragraph(_xml_escape(text), boxed_body_style))
        else:
            story.append(Paragraph("RESPUESTA", label_style))
            for raw_p in text.split("\n"):
                if not raw_p.strip():
                    story.append(Spacer(1, 0.15 * cm))
                    continue
                story.append(Paragraph(format_inline(raw_p), body_style))

            sources = msg.get("sources", [])
            if sources:
                story.append(Spacer(1, 0.2 * cm))
                story.append(Paragraph("FUENTES CONSULTADAS", label_style))
                for s in sources:
                    n = s.get("n", "?")
                    code = s.get("doc_code") or s.get("title", "")
                    page = s.get("page", "")
                    section = s.get("section_short") or s.get("section", "")
                    snippet = (s.get("snippet") or "").strip()
                    chunk_id = s.get("chunk_id", "")
                    score = s.get("score")

                    head_bits = [
                        f'<font face="Courier-Bold" color="{cite_color_hex()}">[{n}]</font>',
                        f"<b>{_xml_escape(str(code))}</b>",
                    ]
                    if section:
                        head_bits.append(
                            f'<font color="#6e737d">{_xml_escape(section)}</font>'
                        )
                    if page:
                        head_bits.append(
                            f'<font color="#6e737d">p. {_xml_escape(str(page))}</font>'
                        )
                    if score is not None:
                        try:
                            head_bits.append(
                                f'<font color="#6e737d" size="9">'
                                f"score {float(score):.2f}</font>"
                            )
                        except (TypeError, ValueError):
                            pass
                    head_html = "  ·  ".join(head_bits)
                    story.append(Paragraph(head_html, source_head_style))

                    if snippet:
                        story.append(Paragraph(f"“{_xml_escape(snippet)}”", snippet_style))
                    if chunk_id:
                        story.append(
                            Paragraph(
                                f"chunk_id: {_xml_escape(str(chunk_id))}",
                                chunkid_style,
                            )
                        )
                story.append(Spacer(1, 0.3 * cm))

    def _on_page(canvas: Any, doc_: Any) -> None:
        canvas.saveState()
        canvas.setFont(f"{base_font}-Oblique", 9)
        canvas.setFillColor(muted_col)
        page_num = canvas.getPageNumber()
        footer_text = (
            f"Generado por LocalDocsAI  ·  página {page_num}"
        )
        canvas.drawCentredString(A4[0] / 2.0, 1.2 * cm, footer_text)
        canvas.restoreState()

    pdf = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=2.2 * cm,
        rightMargin=2.2 * cm,
        topMargin=2.0 * cm,
        bottomMargin=2.2 * cm,
        title=title,
        author="LocalDocsAI",
    )
    pdf.build(story, onFirstPage=_on_page, onLaterPages=_on_page)


def _xml_escape(text: str) -> str:
    return (
        text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    )


def write_conversation_docx(
    messages: list[dict[str, Any]],
    output_path: Path,
    title: str = "Conversación",
    *,
    model_name: str = "",
    app_version: str = "v0.1 · OFFLINE",
) -> None:
    """Export *messages* as a richly formatted Word document.

    Each question/answer pair is rendered as a numbered section with:
        - a Heading 2 derived from the question (so it shows up in any TOC),
        - a labelled PREGUNTA block on a tinted background,
        - the assistant's RESPUESTA with inline [N] citation chips in amber,
        - a FUENTES CONSULTADAS list where every source shows its document
          code, page, section, snippet excerpt, and chunk_id.

    A metadata block at the top records the generation date, model used,
    number of sources consulted, and app version. The footer carries the
    page number and the LocalDocsAI signature.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    doc = Document()

    # ------------------------------------------------------------------
    # Document-wide styles
    # ------------------------------------------------------------------
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(*_TEXT_HEX)

    # Page margins — slightly tighter than Word defaults for a denser look
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ------------------------------------------------------------------
    # Title block
    # ------------------------------------------------------------------
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in heading.runs:
        run.font.color.rgb = RGBColor(*_TEXT_HEX)
        run.font.size = Pt(20)

    subtitle = doc.add_paragraph()
    s_run = subtitle.add_run("Conversación exportada de LocalDocsAI")
    s_run.italic = True
    s_run.font.color.rgb = RGBColor(*_MUTED_HEX)
    s_run.font.size = Pt(10)

    # ------------------------------------------------------------------
    # Metadata block
    # ------------------------------------------------------------------
    now = datetime.now()
    total_sources = sum(len(m.get("sources", [])) for m in messages if m.get("role") == "assistant")
    questions = sum(1 for m in messages if m.get("role") == "user")

    meta_rows = [
        ("Fecha", now.strftime("%d/%m/%Y · %H:%M")),
        ("Aplicación", f"LocalDocsAI  {app_version}"),
        ("Modelo", model_name or "—"),
        ("Preguntas", str(questions)),
        ("Fuentes citadas", str(total_sources)),
    ]
    _add_metadata_table(doc, meta_rows)

    # Thin divider rule before the conversation body
    _add_horizontal_rule(doc)

    # ------------------------------------------------------------------
    # Conversation pairs
    # ------------------------------------------------------------------
    question_index = 0
    pending_question: str | None = None

    for msg in messages:
        role = msg.get("role", "")
        text = msg.get("text", "")

        if role == "user":
            question_index += 1
            pending_question = text
            short = text.strip().split("\n")[0][:80] or "Pregunta"
            qh = doc.add_heading(f"{question_index}.  {short}", level=2)
            for run in qh.runs:
                run.font.color.rgb = RGBColor(*_TEXT_HEX)
                run.font.size = Pt(14)

            _add_label(doc, "PREGUNTA")
            _add_boxed_paragraph(doc, text)

        else:  # assistant
            _add_label(doc, "RESPUESTA")
            _add_response_paragraph(doc, text)

            sources = msg.get("sources", [])
            if sources:
                _add_label(doc, "FUENTES CONSULTADAS")
                _add_sources_block(doc, sources)

            # Trailing breathing room before the next pair
            doc.add_paragraph()
            pending_question = None

    # ------------------------------------------------------------------
    # Footer with page numbers + branding
    # ------------------------------------------------------------------
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.text = ""  # clear default
    _add_run(fp, "Generado por LocalDocsAI  ·  ", italic=True, size=9, color=_MUTED_HEX)
    _add_page_field(fp)
    _add_run(fp, "  de  ", italic=True, size=9, color=_MUTED_HEX)
    _add_page_field(fp, total=True)

    # Force runtime field update on Word open so {PAGE} / {NUMPAGES} render
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        update_fields = settings.makeelement(qn("w:updateFields"), {qn("w:val"): "true"})
        settings.append(update_fields)

    doc.save(str(output_path))


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------


def _add_run(
    paragraph: Any,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    size: int | None = None,
    color: tuple[int, int, int] | None = None,
    mono: bool = False,
) -> Any:
    from docx.shared import Pt, RGBColor

    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor(*color)
    if mono:
        run.font.name = "Consolas"
    return run


def _add_label(doc: Any, label: str) -> None:
    """Small uppercase label above a content block (PREGUNTA, RESPUESTA, etc.)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = 0
    p.paragraph_format.space_after = 0
    _add_run(
        p, label, bold=True, size=8, color=_MUTED_HEX,
    )


def _shade_paragraph(paragraph: Any, hex_color: tuple[int, int, int]) -> None:
    """Add a background-color shading to a paragraph (uses raw XML)."""
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml

    fill = f"{hex_color[0]:02X}{hex_color[1]:02X}{hex_color[2]:02X}"
    shd = parse_xml(
        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{fill}"/>'
    )
    paragraph._p.get_or_add_pPr().append(shd)


def _add_boxed_paragraph(doc: Any, text: str) -> None:
    """Question text in a tinted box for visual separation."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = 0
    p.paragraph_format.right_indent = 0
    p.paragraph_format.space_before = 2
    p.paragraph_format.space_after = 12
    _shade_paragraph(p, _PANEL_HEX)
    _add_run(p, text)


def _add_response_paragraph(doc: Any, text: str) -> None:
    """Render the assistant response, turning [N] tokens into amber bold chips.

    Splits the text into runs around each [N] match so the markers can carry
    different styling (bold, color, slightly smaller) than the body text.
    """
    # Split by paragraphs so blank lines in the original response are
    # preserved as real paragraph breaks in the docx.
    for raw_paragraph in text.split("\n"):
        if not raw_paragraph.strip():
            doc.add_paragraph()
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = 4
        last = 0
        for m in _CITE_RE.finditer(raw_paragraph):
            if m.start() > last:
                _add_run(p, raw_paragraph[last : m.start()])
            _add_run(
                p,
                m.group(0),
                bold=True,
                color=_CITE_HEX,
                size=10,
                mono=True,
            )
            last = m.end()
        if last < len(raw_paragraph):
            _add_run(p, raw_paragraph[last:])


def _add_sources_block(doc: Any, sources: list[dict[str, Any]]) -> None:
    """For each source: [N] doc_code · section · p.X, snippet, chunk_id."""
    from docx.shared import Cm

    for s in sources:
        n = s.get("n", "?")
        code = s.get("doc_code") or s.get("title", "")
        page = s.get("page", "")
        section = s.get("section_short") or s.get("section", "")
        snippet = (s.get("snippet") or "").strip()
        chunk_id = s.get("chunk_id", "")
        score = s.get("score")

        # Header line: [N]  CODE   ·   section   ·   p.X   ·   score%
        head = doc.add_paragraph()
        head.paragraph_format.space_before = 4
        head.paragraph_format.space_after = 0
        head.paragraph_format.left_indent = Cm(0.4)

        _add_run(head, f"[{n}]  ", bold=True, color=_CITE_HEX, mono=True)
        _add_run(head, code, bold=True)
        if section:
            _add_run(head, "  ·  ", color=_MUTED_HEX)
            _add_run(head, section, color=_MUTED_HEX)
        if page:
            _add_run(head, "  ·  ", color=_MUTED_HEX)
            _add_run(head, f"p. {page}", color=_MUTED_HEX)
        if score is not None:
            try:
                _add_run(head, "  ·  ", color=_MUTED_HEX)
                _add_run(head, f"score {float(score):.2f}", color=_MUTED_HEX, size=9)
            except (TypeError, ValueError):
                pass

        # Snippet (italic, indented further)
        if snippet:
            sp = doc.add_paragraph()
            sp.paragraph_format.left_indent = Cm(1.0)
            sp.paragraph_format.right_indent = Cm(0.4)
            sp.paragraph_format.space_before = 0
            sp.paragraph_format.space_after = 0
            _add_run(sp, f"“{snippet}”", italic=True, size=10, color=_MUTED_HEX)

        # chunk_id (small, mono)
        if chunk_id:
            cp = doc.add_paragraph()
            cp.paragraph_format.left_indent = Cm(1.0)
            cp.paragraph_format.space_before = 0
            cp.paragraph_format.space_after = 6
            _add_run(cp, "chunk_id: ", size=8, color=_MUTED_HEX, mono=True)
            _add_run(cp, str(chunk_id), size=8, color=_MUTED_HEX, mono=True)


def _add_metadata_table(doc: Any, rows: list[tuple[str, str]]) -> None:
    """Two-column key/value table with subtle styling."""
    from docx.shared import Cm

    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    table.columns[0].width = Cm(4.5)
    table.columns[1].width = Cm(12.0)

    for i, (k, v) in enumerate(rows):
        cells = table.rows[i].cells
        cells[0].width = Cm(4.5)
        cells[1].width = Cm(12.0)

        kp = cells[0].paragraphs[0]
        _add_run(kp, k.upper(), bold=True, size=8, color=_MUTED_HEX)
        vp = cells[1].paragraphs[0]
        _add_run(vp, v, size=10, color=_TEXT_HEX)

        for cell in cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = 2
                paragraph.paragraph_format.space_after = 2

    # Light bottom border on each row (key/value separation only)
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml

    for row in table.rows:
        for cell in row.cells:
            border_xml = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="DDE2EA"/>'
                "</w:tcBorders>"
            )
            cell._tc.get_or_add_tcPr().append(border_xml)


def _add_horizontal_rule(doc: Any) -> None:
    """Insert a thin horizontal divider between sections."""
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml

    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    border = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '<w:bottom w:val="single" w:sz="6" w:space="1" w:color="DDE2EA"/>'
        "</w:pBdr>"
    )
    pPr.append(border)


def _add_page_field(paragraph: Any, *, total: bool = False) -> None:
    """Insert a {PAGE} or {NUMPAGES} field into *paragraph*."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    fld = OxmlElement("w:fldChar")
    fld.set(qn("w:fldCharType"), "begin")
    run._r.append(fld)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " NUMPAGES " if total else " PAGE "
    run._r.append(instr)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)

    # Small muted style for the page number
    from docx.shared import Pt, RGBColor

    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(*_MUTED_HEX)
    run.italic = True
